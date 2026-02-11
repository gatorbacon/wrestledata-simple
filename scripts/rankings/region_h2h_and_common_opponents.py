#!/usr/bin/env python3
"""
Report region head-to-head wins and region common-opponent wins for a chosen team.

Reads data/team_lists/hs_ky_<gender>/teams.json for region and region teams,
then mt/rankings_data/hs_ky_<gender>/<season>/relationships_<weight>.json
to list each of the team's wrestlers (by weight) with:
  - Region head-to-head wins (opponent, team, date, result)
  - Region common-opponent wins (common opponent, our match, region opponent's match)

Usage:
  python region_h2h_and_common_opponents.py --gender boys --season 2026
  (prompts for team name at runtime)
"""

import argparse
import json
from pathlib import Path
from typing import Dict, List, Set, Tuple, Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def load_teams(gender: str, data_root: Path) -> List[Dict]:
    """Load team list from data/team_lists/hs_ky_<gender>/teams.json."""
    path = data_root / "data" / "team_lists" / f"hs_ky_{gender}" / "teams.json"
    if not path.exists():
        raise FileNotFoundError(f"Team list not found: {path}")
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def find_team_region(teams: List[Dict], team_name: str) -> Tuple[Optional[str], Set[str], Optional[str]]:
    """
    Find the given team and return (region, set of all team names in that region, canonical_team_name).
    Team name is matched case-insensitively; first exact match wins, else first substring match.
    canonical_team_name is the exact name from teams.json (for matching wrestler['team']).
    """
    team_name_clean = team_name.strip()
    team_name_lower = team_name_clean.lower()
    match = None
    for t in teams:
        n = t.get("name") or ""
        if n.strip().lower() == team_name_lower:
            match = t
            break
    if not match:
        for t in teams:
            n = t.get("name") or ""
            if team_name_lower in n.strip().lower():
                match = t
                break
    if not match:
        return None, set(), None
    region = match.get("region")
    if region is None or (isinstance(region, str) and region.strip() == ""):
        return None, set(), None
    region_str = str(region).strip()
    region_teams = {t.get("name", "").strip() for t in teams if str(t.get("region", "")).strip() == region_str}
    canonical_name = (match.get("name") or "").strip()
    return region_str, region_teams, canonical_name


def load_relationships_for_season(season: int, gender: str, data_root: Path) -> Dict[str, Dict]:
    """
    Load all relationships_<weight>.json for the given season/gender.
    Returns dict: weight_class -> { wrestlers, direct_relationships, common_opponent_relationships }.
    Keys in relationship dicts are stored as "id1_id2" in JSON.
    """
    path = data_root / "mt" / "rankings_data" / f"hs_ky_{gender}" / str(season)
    if not path.exists():
        raise FileNotFoundError(f"Rankings data not found: {path}")
    out = {}
    for rel_file in sorted(path.glob("relationships_*.json")):
        wc = rel_file.stem.replace("relationships_", "")
        with rel_file.open("r", encoding="utf-8") as f:
            out[wc] = json.load(f)
    return out


def load_weight_class_data(season: int, gender: str, data_root: Path) -> Dict[str, Dict]:
    """
    Load all weight_class_<weight>.json for the given season/gender.
    Returns dict: weight_class -> { wrestlers, matches }.
    Used to scan ALL matches across ALL weights for cross-weight H2H and common opponents.
    """
    path = data_root / "mt" / "rankings_data" / f"hs_ky_{gender}" / str(season)
    if not path.exists():
        return {}
    out = {}
    for wc_file in sorted(path.glob("weight_class_*.json")):
        wc = wc_file.stem.replace("weight_class_", "")
        with wc_file.open("r", encoding="utf-8") as f:
            out[wc] = json.load(f)
    return out


def _skip_match_for_result(match: Dict) -> bool:
    """True if match should be skipped for H2H/CO (NC, MFF, injury, etc.)."""
    result_str = str(match.get("result", "") or "").lower().strip()
    if not result_str:
        return False
    if result_str == "nc" or "no contest" in result_str:
        return True
    if "mffl" in result_str or "m. for." in result_str or "medical forfeit" in result_str:
        return True
    if "inj" in result_str or "injury" in result_str:
        return True
    return False


def build_all_wrestlers_lookup(relationships_by_weight: Dict[str, Dict]) -> Dict[str, Dict]:
    """Build id -> { name, team } from all weight classes (for common-opponent team lookup)."""
    lookup = {}
    for data in relationships_by_weight.values():
        for wid, w in data.get("wrestlers", {}).items():
            if wid not in lookup:
                lookup[wid] = {"name": (w.get("name") or "").strip(), "team": (w.get("team") or "").strip()}
    return lookup


def build_all_wrestlers_from_weight_data(weight_data_by_weight: Dict[str, Dict]) -> Dict[str, Dict]:
    """Build id -> { name, team } from all weight class files (includes everyone in matches)."""
    lookup = {}
    for data in weight_data_by_weight.values():
        for wid, w in data.get("wrestlers", {}).items():
            if wid not in lookup:
                lookup[wid] = {"name": (w.get("name") or "").strip(), "team": (w.get("team") or "").strip()}
    return lookup


def load_career_lookup(data_root: Path, careers_dir: Optional[Path] = None) -> Dict[str, str]:
    """
    Load all career files and build season_wrestler_id -> career_id.
    Same linkage used by wrestler profile JSON (data/careers/career_*.json).
    """
    base = careers_dir or (data_root / "data" / "careers")
    if not base.exists():
        return {}
    lookup: Dict[str, str] = {}
    for career_file in base.glob("career_*.json"):
        try:
            with career_file.open("r", encoding="utf-8") as f:
                career_data = json.load(f)
            career_id = career_data.get("career_id")
            seasons = career_data.get("seasons", {})
            if career_id and isinstance(seasons, dict):
                for season_wrestler_id in seasons.values():
                    if season_wrestler_id:
                        lookup[str(season_wrestler_id)] = career_id
        except Exception:
            continue
    return lookup


def load_career(career_id: str, data_root: Path, careers_dir: Optional[Path] = None) -> Optional[Dict]:
    """Load a single career JSON (seasons: { year -> season_wrestler_id })."""
    base = careers_dir or (data_root / "data" / "careers")
    path = base / f"{career_id}.json"
    if not path.exists():
        return None
    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def load_last_year_accomplishments(season: int, gender: str, data_root: Path) -> Dict[str, Dict]:
    """
    Load season accomplishments for (season - 1): season_wrestler_id -> { regional_place, state_place }.
    Only includes wrestlers with at least one placement.
    Use with career linkage: current_id -> career -> seasons[last_year] -> this dict.
    """
    path = data_root / "data" / "season_accomplishments" / gender / str(season - 1) / "season_accomplishments.json"
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)
    by_id: Dict[str, Dict] = {}
    for w in data.get("wrestlers", []):
        wid = w.get("season_wrestler_id")
        rp = w.get("regional_place")
        sp = w.get("state_place")
        if rp is None and sp is None:
            continue
        if wid:
            by_id[wid] = {"regional_place": rp, "state_place": sp}
    return by_id


def format_place(n: Optional[int]) -> str:
    """Format 1 -> '1st Place', 2 -> '2nd Place', etc."""
    if n is None:
        return ""
    sufs = {1: "st", 2: "nd", 3: "rd"}
    suffix = sufs.get(n) or "th"
    return f"{n}{suffix} Place"


def _get_last_year_acc(
    w: Dict,
    last_year: int,
    last_year_by_id: Dict[str, Dict],
    career_lookup: Dict[str, str],
    data_root: Path,
) -> Dict:
    """Resolve last-year accomplishments via career linkage. Returns dict with regional_place, state_place or empty."""
    acc = last_year_by_id.get(w["wrestler_id"])
    if acc is None and career_lookup:
        career_id = career_lookup.get(w["wrestler_id"])
        if career_id:
            career = load_career(career_id, data_root)
            if career:
                last_year_id = career.get("seasons", {}).get(str(last_year))
                if last_year_id:
                    acc = last_year_by_id.get(last_year_id)
    return acc or {}


def write_report_docx(
    output_path: Path,
    team_name: str,
    region: str,
    region_teams: Set[str],
    by_weight: Dict[str, List[Dict]],
    *,
    last_year: int,
    last_year_by_id: Dict[str, Dict],
    career_lookup: Dict[str, str],
    data_root: Path,
    all_wrestlers: Dict[str, Dict],
) -> None:
    """Write the report to a Word document. Left-aligned, single spacing, no indentation; names bold, weights blue. Skips wrestlers with no H2H and no common opponents."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for Word output. Install with: pip install python-docx")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Consolas"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.0

    def set_para_format(p):
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    def add_para(text: str = "", bold: bool = False):
        p = doc.add_paragraph()
        set_para_format(p)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.name = "Consolas"
            run.font.size = Pt(11)
        return p

    def add_para_mixed(runs: List[Tuple[str, bool]]):
        p = doc.add_paragraph()
        set_para_format(p)
        for text, bold in runs:
            if not text:
                continue
            run = p.add_run(text)
            run.bold = bold
            run.font.name = "Consolas"
            run.font.size = Pt(11)
        return p

    def add_para_weight_header(wc: str):
        """Add '=== Weight N ===' with N in blue."""
        p = doc.add_paragraph()
        set_para_format(p)
        r1 = p.add_run("=== Weight ")
        r1.font.name = "Consolas"
        r1.font.size = Pt(11)
        r1.bold = True
        r2 = p.add_run(wc)
        r2.font.name = "Consolas"
        r2.font.size = Pt(11)
        r2.bold = True
        r2.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
        r3 = p.add_run(" ===")
        r3.font.name = "Consolas"
        r3.font.size = Pt(11)
        r3.bold = True
        return p

    add_para("Region head-to-head and common-opponent report", bold=False)
    add_para()
    add_para(f"Team: {team_name}  |  Region: {region}", bold=False)
    region_list = ", ".join(sorted(region_teams))
    if len(region_teams) > 20:
        region_list = ", ".join(sorted(region_teams)[:20]) + "..."
    add_para(f"Region teams ({len(region_teams)}): {region_list}", bold=False)
    add_para()

    for wc in sort_weight_keys(list(by_weight.keys())):
        wrestlers = [w for w in by_weight[wc] if w["h2h_wins"] or w["co_wins"]]
        if not wrestlers:
            continue
        add_para_weight_header(wc)
        for w in sorted(wrestlers, key=lambda x: (x["name"].lower(), x["name"])):
            add_para()
            add_para(w["name"], bold=True)
            acc = _get_last_year_acc(w, last_year, last_year_by_id, career_lookup, data_root)
            rp = acc.get("regional_place")
            sp = acc.get("state_place")
            if rp is not None:
                add_para_mixed([(f"{last_year} ", True), ("Regions: ", True), (format_place(rp), False)])
            if sp is not None:
                add_para_mixed([(f"{last_year} ", True), ("States: ", True), (format_place(sp), False)])
            add_para("Head to Head", bold=True)
            for h in w["h2h_wins"]:
                for m in h["matches"]:
                    result = (m.get("result") or "").strip()
                    add_para(f"{h['opponent_name']} ({h['opponent_team']}) {result}", bold=False)
            add_para("Common Opponent", bold=True)
            for c in w["co_wins"]:
                co_team = (all_wrestlers.get(c.get("common_opponent_id") or "") or {}).get("team") or "?"
                ro_name = c.get("region_opponent_name") or "?"
                ro_team = c.get("region_opponent_team") or "?"
                co_name = c.get("common_opponent_name") or "?"
                add_para(f"{ro_name} ({ro_team}) lost to {co_name} ({co_team})", bold=False)
        add_para()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def collect_region_h2h_and_co(
    relationships_by_weight: Dict[str, Dict],
    weight_data_by_weight: Dict[str, Dict],
    canonical_team_name: str,
    region_team_names: Set[str],
    all_wrestlers: Dict[str, Dict],
) -> Dict[str, List[Dict]]:
    """
    For the chosen team, collect per-wrestler region H2H wins and region common-opponent wins
    by scanning ALL matches in ALL weight classes (so cross-weight H2H and CO are included).

    Returns:
      weight_class -> list of { wrestler_id, name, team, weight_class, h2h_wins, co_wins }.
    """
    from collections import defaultdict

    other_region_teams = region_team_names - {canonical_team_name}

    # 1) Our team's wrestlers and their display weight (from relationships, one weight per wrestler)
    our_team_ids: Set[str] = set()
    our_wrestlers_by_id: Dict[str, Dict] = {}  # id -> { name, team, weight_class }
    for weight_class in sort_weight_keys(list(relationships_by_weight.keys())):
        data = relationships_by_weight.get(weight_class, {})
        for w in data.get("wrestlers", {}).values():
            if (w.get("team") or "").strip() != canonical_team_name:
                continue
            wid = w.get("id") or ""
            if not wid or wid in our_wrestlers_by_id:
                continue
            our_team_ids.add(wid)
            our_wrestlers_by_id[wid] = {
                "wrestler_id": wid,
                "name": (w.get("name") or "?").strip(),
                "team": (w.get("team") or "").strip(),
                "weight_class": weight_class,
            }

    # 2) Region opponent ids (any wrestler whose team is in other_region_teams)
    region_opponent_ids: Set[str] = set()
    for wid, winfo in all_wrestlers.items():
        if (winfo.get("team") or "").strip() in other_region_teams:
            region_opponent_ids.add(wid)

    # 3) H2H from ALL matches (all weight files) – include cross-weight wins
    # Dedupe by (our_id, opp_id, date, result) since same match can appear in multiple weight files
    h2h_by_our_id: Dict[str, List[Dict]] = defaultdict(list)
    seen_h2h: Set[Tuple[str, str, str, str]] = set()

    for _wc, data in weight_data_by_weight.items():
        for match in data.get("matches", []):
            if _skip_match_for_result(match):
                continue
            w1 = str(match.get("wrestler1_id") or "")
            w2 = str(match.get("wrestler2_id") or "")
            winner = str(match.get("winner_id") or "")
            if not w1 or not w2 or not winner:
                continue
            # Skip placeholder/invalid ids
            if w1 == "-1" or w2 == "-1" or "OUTSTATE" in w1 or "OUTSTATE" in w2:
                continue
            t1 = (all_wrestlers.get(w1) or {}).get("team") or ""
            t2 = (all_wrestlers.get(w2) or {}).get("team") or ""
            our_id = None
            opp_id = None
            if w1 in our_team_ids and t2 in other_region_teams and winner == w1:
                our_id, opp_id = w1, w2
            elif w2 in our_team_ids and t1 in other_region_teams and winner == w2:
                our_id, opp_id = w2, w1
            if our_id is None or opp_id is None:
                continue
            date = (match.get("date") or "").strip()
            result = (match.get("result") or "").strip()
            key = (our_id, opp_id, date, result)
            if key in seen_h2h:
                continue
            seen_h2h.add(key)
            opp_info = all_wrestlers.get(opp_id) or {}
            h2h_by_our_id[our_id].append({
                "opponent_id": opp_id,
                "opponent_name": (opp_info.get("name") or opp_id).strip(),
                "opponent_team": (opp_info.get("team") or "?").strip(),
                "matches": [{"date": date, "result": result, "event": (match.get("event") or "").strip()}],
            })

    # 4) Common opponents from ALL matches: opponents per wrestler, wins/losses per (wrestler, opponent)
    opponents = defaultdict(set)
    # pair_wins[sorted_pair] = (wins_by_first, wins_by_second)
    pair_wins: Dict[Tuple[str, str], Tuple[int, int]] = defaultdict(lambda: (0, 0))

    for _wc, data in weight_data_by_weight.items():
        for match in data.get("matches", []):
            if _skip_match_for_result(match):
                continue
            w1 = str(match.get("wrestler1_id") or "")
            w2 = str(match.get("wrestler2_id") or "")
            winner = str(match.get("winner_id") or "")
            if not w1 or not w2 or not winner or w1 == "-1" or w2 == "-1":
                continue
            if "OUTSTATE" in w1 or "OUTSTATE" in w2:
                continue
            opponents[w1].add(w2)
            opponents[w2].add(w1)
            pair = tuple(sorted([w1, w2]))
            first, second = pair[0], pair[1]
            wins_first, wins_second = pair_wins[pair]
            if winner == first:
                pair_wins[pair] = (wins_first + 1, wins_second)
            else:
                pair_wins[pair] = (wins_first, wins_second + 1)

    def _wins(wid: str, opp_id: str) -> int:
        p = tuple(sorted([wid, opp_id]))
        a, b = pair_wins.get(p, (0, 0))
        return a if wid == p[0] else b

    def _losses(wid: str, opp_id: str) -> int:
        p = tuple(sorted([wid, opp_id]))
        a, b = pair_wins.get(p, (0, 0))
        return b if wid == p[0] else a

    # 5) CO wins: for each our wrestler and each region opponent, common opponents where we won and they lost
    co_by_our_id: Dict[str, List[Dict]] = defaultdict(list)
    seen_co: Set[Tuple[str, str, str]] = set()

    for our_id in our_team_ids:
        for reg_id in region_opponent_ids:
            common = opponents.get(our_id, set()) & opponents.get(reg_id, set())
            for co_id in common:
                if _wins(our_id, co_id) == 0 or _losses(reg_id, co_id) == 0:
                    continue
                key = (our_id, reg_id, co_id)
                if key in seen_co:
                    continue
                seen_co.add(key)
                reg_info = all_wrestlers.get(reg_id) or {}
                co_info = all_wrestlers.get(co_id) or {}
                co_by_our_id[our_id].append({
                    "region_opponent_id": reg_id,
                    "region_opponent_name": (reg_info.get("name") or reg_id).strip(),
                    "region_opponent_team": (reg_info.get("team") or "?").strip(),
                    "common_opponent_id": co_id,
                    "common_opponent_name": (co_info.get("name") or co_id).strip(),
                })

    # 6) Build result by weight (group our wrestlers by display weight)
    result_by_weight: Dict[str, List[Dict]] = defaultdict(list)
    for wid, info in our_wrestlers_by_id.items():
        wc = info["weight_class"]
        result_by_weight[wc].append({
            "wrestler_id": wid,
            "name": info["name"],
            "team": info["team"],
            "weight_class": wc,
            "h2h_wins": h2h_by_our_id.get(wid, []),
            "co_wins": co_by_our_id.get(wid, []),
        })

    return dict(result_by_weight)


def sort_weight_keys(weight_classes: List[str]) -> List[str]:
    """Sort weight class strings numerically where possible."""
    def key(w):
        try:
            return (0, int(w))
        except ValueError:
            return (1, w)
    return sorted(weight_classes, key=key)


def print_report(
    team_name: str,
    region: str,
    region_teams: Set[str],
    by_weight: Dict[str, List[Dict]],
    *,
    last_year: int,
    last_year_by_id: Dict[str, Dict],
    career_lookup: Dict[str, str],
    data_root: Path,
    all_wrestlers: Dict[str, Dict],
) -> None:
    """Print the report to stdout. Last-year placements resolved via career linkage (season id -> career -> last year season id)."""
    print(f"\nRegion head-to-head and common-opponent report")
    print(f"Team: {team_name}  |  Region: {region}")
    print(f"Region teams ({len(region_teams)}): {', '.join(sorted(region_teams)[:15])}{'...' if len(region_teams) > 15 else ''}")
    print()

    for wc in sort_weight_keys(list(by_weight.keys())):
        wrestlers = [w for w in by_weight[wc] if w["h2h_wins"] or w["co_wins"]]
        if not wrestlers:
            continue
        print(f"=== Weight {wc} ===")
        for w in sorted(wrestlers, key=lambda x: (x["name"].lower(), x["name"])):
            print(f"\n{w['name']}")
            acc = _get_last_year_acc(w, last_year, last_year_by_id, career_lookup, data_root)
            rp = acc.get("regional_place")
            sp = acc.get("state_place")
            if rp is not None:
                print(f"{last_year} Regions: {format_place(rp)}")
            if sp is not None:
                print(f"{last_year} States: {format_place(sp)}")
            if rp is not None or sp is not None:
                print()
            print("Head to Head")
            for h in w["h2h_wins"]:
                for m in h["matches"]:
                    result = (m.get("result") or "").strip()
                    print(f"{h['opponent_name']} ({h['opponent_team']})  {result}")
            print("Common Opponent")
            for c in w["co_wins"]:
                co_team = "?"
                co_id = c.get("common_opponent_id")
                if co_id and all_wrestlers:
                    co_team = (all_wrestlers.get(co_id) or {}).get("team") or "?"
                ro_name = c.get("region_opponent_name") or "?"
                ro_team = c.get("region_opponent_team") or "?"
                co_name = c.get("common_opponent_name") or "?"
                print(f"{ro_name} ({ro_team}) lost to {co_name} ({co_team})")
        print()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Report region head-to-head and common-opponent wins for a team (HS KY)."
    )
    parser.add_argument("--gender", required=True, choices=["boys", "girls"], help="Gender (boys/girls)")
    parser.add_argument("--season", type=int, required=True, help="Season year (e.g. 2026)")
    parser.add_argument("--data-dir", type=Path, default=Path(__file__).resolve().parents[2], help="Repo root (default: two levels up from script)")
    parser.add_argument("-o", "--output", type=Path, default=None, help="Output path for Word document (.docx). If omitted, report is printed to stdout.")
    args = parser.parse_args()
    data_root = args.data_dir

    team_name = input("Team name: ").strip()
    if not team_name:
        print("No team name provided. Exiting.")
        return

    teams = load_teams(args.gender, data_root)
    region, region_team_names, canonical_team_name = find_team_region(teams, team_name)
    if region is None or canonical_team_name is None:
        print(f"Team '{team_name}' not found or has no region assigned.")
        return
    if not region_team_names:
        print(f"Team '{team_name}' has region '{region}' but no other teams in that region found.")
        return

    relationships_by_weight = load_relationships_for_season(args.season, args.gender, data_root)
    if not relationships_by_weight:
        print(f"No relationship data found for season {args.season} ({args.gender}).")
        return

    weight_data_by_weight = load_weight_class_data(args.season, args.gender, data_root)
    if not weight_data_by_weight:
        print(f"No weight class data found for season {args.season} ({args.gender}).")
        return

    # Use wrestlers from weight class data so we have team for everyone in matches (cross-weight opponents)
    all_wrestlers = build_all_wrestlers_from_weight_data(weight_data_by_weight)
    # Merge in any from relationships not in weight data (e.g. wrestlers with no matches)
    for data in relationships_by_weight.values():
        for wid, w in data.get("wrestlers", {}).items():
            if wid not in all_wrestlers:
                all_wrestlers[wid] = {"name": (w.get("name") or "").strip(), "team": (w.get("team") or "").strip()}

    by_weight = collect_region_h2h_and_co(
        relationships_by_weight,
        weight_data_by_weight,
        canonical_team_name,
        region_team_names,
        all_wrestlers,
    )
    if not by_weight:
        print(f"No wrestlers found for team '{team_name}' in relationship data.")
        return

    last_year = args.season - 1
    last_year_by_id = load_last_year_accomplishments(args.season, args.gender, data_root)
    career_lookup = load_career_lookup(data_root)

    report_kw = dict(
        last_year=last_year,
        last_year_by_id=last_year_by_id,
        career_lookup=career_lookup,
        data_root=data_root,
        all_wrestlers=all_wrestlers,
    )
    if args.output is not None:
        if not DOCX_AVAILABLE:
            print("Error: Word output requires python-docx. Install with: pip install python-docx")
            return
        out_path = args.output
        if out_path.suffix.lower() != ".docx":
            out_path = out_path.with_suffix(".docx")
        write_report_docx(
            out_path,
            canonical_team_name,
            region,
            region_team_names,
            by_weight,
            **report_kw,
        )
        print(f"Report written to {out_path}")
    else:
        print_report(
            canonical_team_name,
            region,
            region_team_names,
            by_weight,
            **report_kw,
        )


if __name__ == "__main__":
    main()
